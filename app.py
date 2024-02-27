from flask import Flask, request, render_template_string, redirect, url_for, render_template
import docx
import translator
import summarise
import question
from flask import make_response
from gtts import gTTS
import os
from playsound import playsound

global language 

app = Flask(__name__)

# HTML template for the upload form with some basic styling
UPLOAD_FORM = '''
<!doctype html>
<html>
<head>
  <title>Upload File</title>
  <style>
    body { font-family: Arial, sans-serif; margin: 20px; }
    h2 { color: #0053bc; }
    .upload-form { margin-top: 20px; }
    input[type=file], input[type=submit], button { margin-top: 10px; }
    input[type=submit], button { 
      background-color: #0053bc; color: white; padding: 10px 20px; 
      border: none; border-radius: 4px; cursor: pointer; }
    input[type=submit]:hover, button:hover { background-color: #0053bc; }
    .answer { margin-top: 20px; }
  </style>
</head>
<body>
  <h2>Upload Text or DOCX File</h2>
  <form method=post enctype=multipart/form-data class="upload-form">
    <input type=file name=textfile>
    <input type=submit value=Upload>
  </form>
</body>
</html>
'''

# HTML template for displaying the text with "Translate", "Summarise", and "Ask a Question" buttons
DISPLAY_TEXT_TEMPLATE = '''
<!doctype html>
<html>
<head>
  <title>File Content</title>
  <style>
    body { font-family: Arial, sans-serif; margin: 20px; }
    .container { display: flex; flex-direction: column; max-width: 100%; }
    .text { width: 100%; }
    .text pre, .translation pre { 
      white-space: pre-wrap; 
      word-wrap: break-word; 
      overflow: auto; 
      max-height: 400px; 
      border: 1px solid #ccc; 
      padding: 10px; 
      border-radius: 5px; 
      background-color: #f8f8f8; 
    }
    input[type=submit], button {
      background-color: #0053bc; color: white; padding: 10px 20px; 
      margin-top: 20px; border: none; border-radius: 4px; cursor: pointer;
    }
    input[type=submit]:hover, button:hover { background-color: #0053bc; }
    .answer { margin-top: 20px; }
  </style>
</head>
<body>
  <div class="container">
    <div class="text">
      <h2>Trial Document:</h2>
      <pre>{{ content }}</pre>
      <form action="/translate" method="post">
        <input id = "doc" type="hidden" name="content" value="{{ content }}">
        <input type="submit" value="Translate">
      </form>
      <form action="/summarise" method="post">
        <input type="hidden" name="content" value="{{ content }}">
        <input type="submit" value="Summarise">
      </form>
      <form action="/question" method="post">
        <input id = "question", type="hidden" name="content" value="{{ content }}">
        <input type="text" name="question" placeholder="Ask a question">
        <input type="submit" value="Ask">
      </form>
      <button onclick="convertToSpeech()">Speak</button>
      <button onclick="">Sign</button>
      {% if answer %}
      <div class="answer">
        <strong>Answer:</strong>
        <p>{{ answer }}</p>
      </div>
      {% endif %}
    </div>
  </div>
  <a href="/">Upload another file</a>
</body>


  <script>
    function convertToSpeech() {
      var text = document.getElementById('question').value;
      fetch('/text_to_speech', {
        method: 'POST',
        headers: {
          'Content-Type': 'application/x-www-form-urlencoded',
        },
        body: 'text=' + encodeURIComponent(text)
      })
      .then(response => response.blob())
      .then(blob => {
        var url = URL.createObjectURL(blob);
        var audio = new Audio(url);
        audio.play();
      })
      .catch(error => console.error('Error:', error));
    }
  </script>
</html>
'''



@app.route('/upload_file', methods=['GET', 'POST'])
def upload_file():

    global language 
    language = "English"
    if request.method == 'POST':
        if 'textfile' not in request.files:
            return 'No file part'
        file = request.files['textfile']
        if file.filename == '':
            return 'No selected file'
        if file:
            if file.filename.endswith('.docx'):
                doc = docx.Document(file)
                content = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
            elif file.filename.endswith('.txt'):
                content = file.read().decode('utf-8')
            else:
                return "This application supports .txt and .docx files only."
            return render_template_string(DISPLAY_TEXT_TEMPLATE, content=content)
    return UPLOAD_FORM

@app.route('/process_unique_id', methods=['POST'])
def process_unique_id():
    unique_id = request.form['uniqueID']
    
    # Based on the uniqueID, determine which HTML file to serve
    if unique_id == '123':  # Example condition
        return render_template('patient_lang.html')
    elif unique_id == '456':  # Another example condition
        return render_template('researches.html')
    else:
        # If uniqueID doesn't match any condition, you can redirect to a default page
        return render_template('default.html')

@app.route('/translate', methods=['POST'])
def translate_text():
    original_content = request.form['content']
    # Placeholder for translation logic
    translated_content = "Simulated translation of the uploaded document."
    translated_content = translator.translate("spanish",original_content)
    # Render the template with the translated content
    return render_template_string(DISPLAY_TEXT_TEMPLATE, content=translated_content)

@app.route('/summarise', methods=['POST'])
def summarise_text():
    original_content = request.form['content']
    # Placeholder for translation logic
    translated_content = "Simulated translation of the uploaded document."
    translated_content = summarise.summarise(language,original_content)
    # Render the template with the translated content
    return render_template_string(DISPLAY_TEXT_TEMPLATE, content=translated_content)

@app.route('/question', methods=['POST'])
def question_text():
    global language
    language="spanish"
    content = request.form['content']
    user_question = request.form['question']
    answer = question.question(language, content, user_question )
    return render_template_string(DISPLAY_TEXT_TEMPLATE, content=content, answer=answer)



@app.route('/display_patient', methods=['GET', 'POST'])
def display_patient():
    global language
    language = request.form['language']
    if request.method == 'POST':


        try:
            # Open the .docx file from the local file system
            doc = docx.Document("ICF_001.docx")
            content = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
        except IOError:
            content = "Failed to load the document."

    # Render the template with the content of the document
    #return render_template('display_text.html', content=content, answer=None)
    response = make_response(render_template_string(DISPLAY_TEXT_TEMPLATE, content=content))
    response.set_cookie('language', language, max_age=60*60*24*30)  # Expires in 30 days
    return response


@app.route('/', methods=['GET', 'POST'])
def ask_id():

    # Render the template with the content of the document
    return render_template('id.html', answer=None)
    #return render_template_string(DISPLAY_TEXT_TEMPLATE, content=content)


@app.route('/text_to_speech', methods=['GET', 'POST'])
def text_to_speech():
    global language
    text = request.form['text']
    if language=="spanish":
        lang = "es"
    else :
        lang = "en"
    tts = gTTS(text=text, lang=lang, slow=False)
    filename = "speech.mp3"
    tts.save(filename)
    playsound(filename)
    #return filename

if __name__ == '__main__':
    app.run(debug=True)

